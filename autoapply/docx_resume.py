"""Adaptação do currículo editando o .docx base, em vez de gerar um do zero.

O fluxo antigo montava um JSON Resume e renderizava num template HTML próprio: saía
uma página tipograficamente pobre e com a voz de LLM. Aqui o ponto de partida é o
currículo real do usuário, e a adaptação é cirúrgica — trocar a ordem e a ênfase de
alguns bullets para a vaga, sem tocar em estilo, seções ou fatos.

Editor primário: o Claude Code CLI que roda na VPS, com as ferramentas de escrita
dele sobre uma cópia isolada do arquivo. Se ele falhar ou estragar a formatação, cai
para o Gemini, que devolve substituições pontuais aplicadas de forma determinística.
"""
from __future__ import annotations

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)

CLAUDE_BIN = os.environ.get("CLAUDE_BIN", "/opt/claude/versions/2.1.220")
CLAUDE_MODEL = os.environ.get("CLAUDE_MODEL", "opus")
CLAUDE_FALLBACK = os.environ.get("CLAUDE_FALLBACK_MODEL", "sonnet")
CLAUDE_TIMEOUT = int(os.environ.get("CLAUDE_TIMEOUT", "600"))

# Marcas que denunciam texto de LLM. A lista é de combate: entra o que aparece de
# fato nas saídas, não o que soa suspeito em tese.
TRAVESSOES = "—–―"
CLICHES = (
    "robusto", "robusta", "sólido", "sólida", "abrangente", "aproveitando",
    "impulsionar", "alavancar", "de ponta", "escalável e eficiente",
    "com o objetivo de otimizar", "desempenhou um papel", "não apenas",
    "mas também", "vale ressaltar", "é importante notar", "em suma",
)

CLICHES_EN = (
    "robust", "seamless", "cutting-edge", "leverage", "leveraging", "spearheaded",
    "delve", "tapestry", "underscore", "it's worth noting", "furthermore",
    "in today's fast-paced", "passionate about", "proven track record",
    "results-driven", "synergy", "holistic",
)

REGRAS_DE_VOZ_EN = f"""WRITING RULES (mandatory):
1. NEVER use an em dash (—) or en dash (–) in prose. Use a comma, a period or a
   colon. The date ranges already in the document stay exactly as they are.
2. Third person, past tense, matching the original: "Developed", "Implemented",
   "Built". No first person, no "I".
3. Banned: {", ".join(CLICHES_EN[:12])}. No inflated adjectives.
4. Short, concrete sentences. Prefer the number and the technology over praise:
   "Cut inspection time using YOLOv5" beats "Developed a robust, scalable solution".
5. Invent nothing. No company, technology, number, certification or date that is not
   already in the résumé. Reorder, rewrite and trim; never create.
6. Keep each entry close to its original length.
7. Natural US English. Do not translate proper nouns, course names or company names
   that have no English equivalent, and keep university names in Portuguese with a
   short English gloss only when it helps.
"""

REGRAS_DE_VOZ = f"""REGRAS DE ESCRITA (obrigatórias):
1. NUNCA use travessão (—), meia-risca (–) ou barra vertical decorativa. Se precisar
   separar ideias, use vírgula, ponto ou dois-pontos.
2. Escreva como o currículo original já escreve: terceira pessoa, verbo no passado
   ("Desenvolveu", "Implementou", "Estruturou"). Não mude a pessoa verbal.
3. Proibido: {", ".join(CLICHES[:12])}. Nada de adjetivo inflado.
4. Frase curta e concreta. Prefira o número e a tecnologia ao elogio: "Reduziu o tempo
   de inspeção usando YOLOv5" vale mais que "Desenvolveu solução robusta e escalável".
5. Não invente nada. Nenhuma empresa, tecnologia, número, certificação ou período que
   já não esteja no currículo. Você pode reordenar, reescrever e cortar; nunca criar.
6. Mantenha o tamanho parecido com o original. Não infle nem resuma demais.
7. Português do Brasil, sem gerundismo e sem voz passiva desnecessária.
"""


@dataclass
class ResultadoAdaptacao:
    caminho: Path
    editor: str                       # "claude" | "gemini"
    edicoes: int = 0
    idioma: str = "pt"                # idioma em que o currículo foi entregue
    resumo: str = ""
    avisos: list[str] = field(default_factory=list)


# --------------------------------------------------------------------------- leitura
def esboco(caminho: Path) -> list[dict]:
    """Índice, estilo e texto de cada parágrafo não vazio — o mapa que o LLM edita."""
    import docx

    doc = docx.Document(str(caminho))
    return [
        {"i": i, "estilo": p.style.name, "texto": p.text}
        for i, p in enumerate(doc.paragraphs)
        if p.text.strip()
    ]


def _texto_completo(caminho: Path) -> str:
    import docx

    return "\n".join(p.text for p in docx.Document(str(caminho)).paragraphs)


# ------------------------------------------------------------------------- aplicação
def aplicar_edicoes(origem: Path, destino: Path, edicoes: dict[int, str]) -> int:
    """Grava as substituições preservando a formatação do parágrafo.

    O texto entra no primeiro run e os demais são esvaziados, em vez de recriar o
    parágrafo: assim fonte, tamanho, negrito, recuo e o marcador da lista continuam
    exatamente como estavam no arquivo do usuário.
    """
    import docx

    doc = docx.Document(str(origem))
    aplicadas = 0
    for indice, novo in edicoes.items():
        if not (0 <= indice < len(doc.paragraphs)):
            log.warning("Índice de parágrafo fora do documento: %s", indice)
            continue
        par = doc.paragraphs[indice]
        if not par.runs:
            continue
        par.runs[0].text = novo
        for run in par.runs[1:]:
            run.text = ""
        aplicadas += 1
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destino))
    return aplicadas


# ------------------------------------------------------------------------- validação
def validar(origem: Path, destino: Path, traduzido: bool = False) -> list[str]:
    """A adaptação preservou o documento? Devolve a lista de problemas.

    É a rede de segurança que permite dar liberdade de edição ao agente: se ele
    reescrever o arquivo de um jeito que perca seções ou estilos, isto detecta e o
    chamador cai para o caminho determinístico.

    Com `traduzido`, o texto dos títulos pode mudar ("Educação" -> "Education") — a
    checagem passa a ser sobre a quantidade e a posição das seções, não sobre as
    palavras. Sem isso, toda tradução seria reprovada como documento quebrado.
    """
    import docx

    problemas: list[str] = []
    a, b = docx.Document(str(origem)), docx.Document(str(destino))

    est_a = sorted(p.style.name for p in a.paragraphs if p.text.strip())
    est_b = sorted(p.style.name for p in b.paragraphs if p.text.strip())
    if est_a != est_b:
        problemas.append(
            f"conjunto de estilos mudou ({len(est_a)} parágrafos viraram {len(est_b)})")

    tit_a = [p.text.strip() for p in a.paragraphs if p.style.name.startswith("Heading")]
    tit_b = [p.text.strip() for p in b.paragraphs if p.style.name.startswith("Heading")]
    if traduzido:
        if len(tit_a) != len(tit_b):
            problemas.append(f"número de seções mudou: {len(tit_a)} -> {len(tit_b)}")
        vazias = [t for t in tit_b if not t]
        if vazias:
            problemas.append("seção ficou sem título")
    elif tit_a != tit_b:
        problemas.append(f"seções mudaram: {tit_a} -> {tit_b}")

    # Travessão e clichê só contam quando a edição os INTRODUZIU. O currículo base
    # usa meia-risca nas datas ("Jan. 2025 – Atualmente"): é a formatação do usuário,
    # não marca de LLM, e reprovar por isso descartaria adaptação boa.
    texto = _texto_completo(destino)
    antes = {p["i"]: p["texto"] for p in esboco(origem)}
    novos = [p["texto"] for p in esboco(destino) if antes.get(p["i"]) != p["texto"]]
    texto_novo = "\n".join(novos)

    achados = [c for c in TRAVESSOES if c in texto_novo]
    if achados:
        problemas.append(f"travessão introduzido: {achados}")

    minusculo = texto_novo.lower()
    vocabulario = CLICHES_EN if traduzido else CLICHES
    cliches = [c for c in vocabulario if c in minusculo]
    if cliches:
        problemas.append(f"clichê de LLM: {cliches[:5]}")

    # Nome e e-mail são os âncoras da identidade: se sumiram, o arquivo está errado.
    for obrigatorio in ("Rhayron", "@"):
        if obrigatorio not in texto:
            problemas.append(f"perdeu conteúdo essencial: {obrigatorio!r}")

    return problemas


def limpar_texto(texto: str) -> str:
    """Remove as marcas mecânicas de LLM que dá para corrigir sem reescrever.

    A meia-risca entre números fica: é o intervalo de datas do currículo original
    ("Jan. 2025 – Atualmente"). Só o uso como travessão de prosa é convertido.
    """
    texto = re.sub(r"\s*[—―]\s*", ", ", texto)
    texto = re.sub(r"(?<![0-9])\s+–\s+(?![0-9])", ", ", texto)
    texto = re.sub(r"\s+([,.;:])", r"\1", texto)
    return re.sub(r"[ \t]{2,}", " ", texto).strip()


# ------------------------------------------------------------------ editor: claude cli
def _prompt_claude(vaga_txt: str, docx_nome: str, idioma: str) -> str:
    if idioma == "en":
        tarefa = """- A vaga está em INGLÊS, e o currículo está em português. Traduza o
  documento INTEIRO para inglês, incluindo os títulos de seção ("Educação" vira
  "Education", "Habilidades Técnicas" vira "Technical Skills", e assim por diante).
- Traduzir e adaptar são a mesma passagem: ao verter cada bullet, escolha o termo em
  inglês que a vaga usa. Não traduza ao pé da letra o que o mercado escreve de outro
  jeito.
- Aqui você edita TODOS os parágrafos com texto, não só alguns. Mantenha nomes
  próprios, empresas e nomes de curso como estão."""
        regras = REGRAS_DE_VOZ_EN
    else:
        tarefa = """- Faça alterações PONTUAIS: reescreva o texto de alguns bullets de
  experiência e de habilidades para aproximar o vocabulário do currículo ao da vaga,
  promovendo o que é relevante e reduzindo o que não é.
- Altere no máximo 12 parágrafos. Se um parágrafo já serve, deixe como está.
- NÃO altere os títulos de seção."""
        regras = REGRAS_DE_VOZ

    return f"""Você vai adaptar um currículo em .docx para uma vaga específica.

ARQUIVO: {docx_nome} (no diretório atual). Edite-o no lugar, usando python-docx.

VAGA:
{vaga_txt}

O QUE FAZER:
- Leia o .docx e entenda a estrutura. Não mude estilos, ordem das seções, nem o
  número de parágrafos.
{tarefa}
- Ao editar um parágrafo com python-docx, escreva no primeiro run e esvazie os demais
  (`p.runs[0].text = novo` e `run.text = ""` nos outros), para não perder formatação.

{regras}

Ao terminar, escreva um arquivo `resumo.txt` no diretório atual com 3 a 6 linhas em
PORTUGUÊS dizendo o que mudou e por quê, na primeira pessoa e sem jargão.
Não escreva mais nada no stdout além de OK ao final."""


def _editar_com_claude(base: Path, saida: Path, vaga_txt: str,
                       idioma: str = "pt") -> ResultadoAdaptacao:
    if not Path(CLAUDE_BIN).exists():
        raise FileNotFoundError(f"claude não encontrado em {CLAUDE_BIN}")

    with tempfile.TemporaryDirectory() as tmp:
        trabalho = Path(tmp)
        alvo = trabalho / "curriculo.docx"
        shutil.copy2(base, alvo)

        cmd = [
            CLAUDE_BIN, "-p", _prompt_claude(vaga_txt, alvo.name, idioma),
            "--model", CLAUDE_MODEL,
            "--fallback-model", CLAUDE_FALLBACK,
            # Editar .docx exige rodar python-docx, ou seja a ferramenta Bash — que
            # acceptEdits não libera: o agente terminava sem tocar no arquivo. O
            # sandbox aqui é o diretório temporário com uma cópia isolada, dentro do
            # container. Liberar as ferramentas nominalmente em vez de desligar o
            # sistema de permissões, que o CLI recusa quando o processo é root.
            "--permission-mode", "acceptEdits",
            "--allowedTools", "Bash", "Read", "Write", "Edit", "Glob", "Grep",
            "--add-dir", str(trabalho),
        ]
        proc = subprocess.run(cmd, cwd=trabalho, capture_output=True, text=True,
                              timeout=CLAUDE_TIMEOUT)
        if proc.returncode != 0:
            raise RuntimeError(
                f"claude saiu com {proc.returncode}: {(proc.stderr or proc.stdout)[:400]}")

        if alvo.stat().st_size == 0:
            raise RuntimeError("claude deixou o .docx vazio")

        resumo_txt = trabalho / "resumo.txt"
        resumo = resumo_txt.read_text(encoding="utf-8").strip() if resumo_txt.exists() else ""

        antes = {p["i"]: p["texto"] for p in esboco(base)}
        depois = {p["i"]: p["texto"] for p in esboco(alvo)}
        mudados = sum(1 for i, t in depois.items() if antes.get(i) != t)
        if mudados == 0:
            # Sair com sucesso sem editar nada é falha silenciosa, não adaptação.
            raise RuntimeError("claude terminou sem alterar o documento")

        saida.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(alvo, saida)

    return ResultadoAdaptacao(caminho=saida, editor="claude", edicoes=mudados,
                              resumo=limpar_texto(resumo))


# -------------------------------------------------------------------- editor: gemini
def _sistema_gemini(idioma: str) -> str:
    if idioma == "en":
        instrucao = """A vaga está em INGLÊS. Traduza o currículo INTEIRO para inglês e
adapte na mesma passagem: ao verter cada trecho, use o termo que a vaga usa, não a
tradução literal. Inclua TODOS os parágrafos com texto na resposta, inclusive os
títulos de seção ("Educação" -> "Education", "Habilidades Técnicas" -> "Technical
Skills"). Mantenha nomes de pessoa, empresa e curso como estão."""
        regras = REGRAS_DE_VOZ_EN
    else:
        instrucao = """Edite no máximo 12 parágrafos. Nunca edite títulos de seção
(estilo Heading), o nome nem a linha de contato. Se um parágrafo já serve para a
vaga, não o inclua."""
        regras = REGRAS_DE_VOZ

    return f"""Você adapta currículos para vagas específicas, de forma cirúrgica.

Recebe os parágrafos numerados de um currículo e a descrição de uma vaga. Devolve as
substituições que aproximam o currículo da vaga.

{regras}

{instrucao}

Responda um JSON assim, sem nada em volta:
{{"edicoes": {{"15": "novo texto do parágrafo 15", "62": "novo texto do 62"}},
  "resumo": "3 a 6 linhas EM PORTUGUÊS sobre o que mudou e por quê"}}"""


def _editar_com_gemini(llm, base: Path, saida: Path, vaga_txt: str,
                       idioma: str = "pt") -> ResultadoAdaptacao:
    paragrafos = esboco(base)
    mapa = "\n".join(f'{p["i"]} [{p["estilo"]}] {p["texto"]}' for p in paragrafos)
    user = f"VAGA:\n{vaga_txt}\n\nCURRÍCULO (índice [estilo] texto):\n{mapa}"

    data = llm.complete_json(_sistema_gemini(idioma), user)
    brutas = data.get("edicoes") or {}

    # Traduzindo, título de seção e linha de contato também mudam; em adaptação
    # normal eles são intocáveis.
    protegidos: set[int] = set() if idioma == "en" else {
        p["i"] for p in paragrafos if p["estilo"].startswith("Heading") or p["i"] <= 1
    }
    edicoes: dict[int, str] = {}
    for chave, texto in brutas.items():
        try:
            i = int(chave)
        except (TypeError, ValueError):
            continue
        if i in protegidos or not isinstance(texto, str) or not texto.strip():
            continue
        edicoes[i] = limpar_texto(texto)

    n = aplicar_edicoes(base, saida, edicoes)
    return ResultadoAdaptacao(caminho=saida, editor="gemini", edicoes=n, idioma=idioma,
                              resumo=limpar_texto(str(data.get("resumo") or "")))


# ------------------------------------------------------------------------- orquestração
def adaptar(job, base: Path, saida: Path, llm=None) -> ResultadoAdaptacao:
    """Adapta o currículo para a vaga. Claude primeiro, Gemini como rede.

    A validação roda nos dois caminhos: um .docx que perdeu seções ou ganhou
    travessão é tratado como falha, não como resultado aceitável.
    """
    from .idioma import detectar

    idioma = detectar(job.title, job.description)
    traduz = idioma == "en"
    log.info("Vaga detectada em %s; currículo será entregue em %s",
             idioma.upper(), "inglês" if traduz else "português")

    vaga_txt = (f"{job.title} @ {job.company}\n"
                f"Local: {job.location or 'não informado'}\n\n"
                f"{(job.description or '')[:6000]}")

    try:
        r = _editar_com_claude(base, saida, vaga_txt, idioma)
        r.idioma = idioma
        problemas = validar(base, r.caminho, traduzido=traduz)
        if problemas:
            raise RuntimeError(f"claude quebrou a formatação: {problemas}")
        log.info("CV adaptado pelo Claude (%d parágrafos, %s)", r.edicoes, idioma)
        return r
    except Exception as e:  # noqa: BLE001
        log.warning("Claude indisponível ou inválido (%s); caindo para o Gemini",
                    str(e)[:200])

    if llm is None:
        raise RuntimeError("Claude falhou e nenhum LLM de fallback foi fornecido")

    r = _editar_com_gemini(llm, base, saida, vaga_txt, idioma)
    r.avisos = validar(base, r.caminho, traduzido=traduz)
    if r.avisos:
        log.warning("Fallback com ressalvas: %s", r.avisos)
    log.info("CV adaptado pelo Gemini (%d parágrafos, %s)", r.edicoes, idioma)
    return r
